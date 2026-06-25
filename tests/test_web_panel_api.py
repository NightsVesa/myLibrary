import json
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import httpx
from fastapi.testclient import TestClient
from PIL import Image

from storage.note_meta import add_recent, set_favorite, set_tags
from llm.wiki_lint import LintFinding
import web_panel.api as api_module
from web_panel.api import create_app


def _client(notes_dir: Path, token: str = "panel-token", **kwargs) -> TestClient:
    return TestClient(create_app(notes_dir=notes_dir, panel_token=token, **kwargs))


def _auth(token: str = "panel-token") -> dict[str, str]:
    return {"X-Panel-Token": token}


def test_rejects_requests_without_panel_token(tmp_path):
    client = _client(tmp_path)

    response = client.get("/api/search", params={"q": "alpha"})

    assert response.status_code == 401
    assert response.json() == {
        "ok": False,
        "data": None,
        "error": {"code": "unauthorized", "message": "Invalid panel token"},
    }


def test_health_uses_uniform_response_shape(tmp_path):
    client = _client(tmp_path)

    response = client.get("/api/health", headers=_auth())

    assert response.status_code == 200
    assert response.json()["ok"] is True
    assert response.json()["data"]["status"] == "ok"
    assert response.json()["error"] is None


def test_unhandled_api_errors_use_uniform_json_response(monkeypatch, tmp_path):
    def fail_search(*_args, **_kwargs):
        raise RuntimeError("search exploded")

    monkeypatch.setattr(api_module, "search_notes_ranked", fail_search, raising=False)
    client = TestClient(
        create_app(notes_dir=tmp_path, panel_token="panel-token"),
        raise_server_exceptions=False,
    )

    response = client.get(
        "/api/search",
        params={"q": "alpha", "mode": "fulltext"},
        headers=_auth(),
    )

    assert response.status_code == 500
    payload = response.json()
    assert payload["ok"] is False
    assert payload["error"]["code"] == "internal_error"
    assert "search exploded" in payload["error"]["message"]


def test_search_supports_fulltext_recent_favorite_and_tag_modes(tmp_path):
    alpha = tmp_path / "alpha.md"
    beta = tmp_path / "beta.md"
    alpha.write_text("# Alpha\nkeyword body", encoding="utf-8")
    beta.write_text("# Beta\nother body", encoding="utf-8")
    set_favorite(alpha, True, notes_dir=tmp_path)
    set_tags(beta, "project", notes_dir=tmp_path)
    add_recent(beta, notes_dir=tmp_path)
    client = _client(tmp_path)

    fulltext = client.get(
        "/api/search",
        params={"q": "keyword", "mode": "fulltext"},
        headers=_auth(),
    ).json()
    recent = client.get(
        "/api/search",
        params={"mode": "recent"},
        headers=_auth(),
    ).json()
    favorite = client.get(
        "/api/search",
        params={"mode": "favorite"},
        headers=_auth(),
    ).json()
    tagged = client.get(
        "/api/search",
        params={"q": "project", "mode": "tag"},
        headers=_auth(),
    ).json()

    assert [item["name"] for item in fulltext["data"]["results"]] == ["alpha.md"]
    assert fulltext["data"]["results"][0]["favorite"] is True
    assert [item["name"] for item in recent["data"]["results"]] == ["beta.md"]
    assert [item["name"] for item in favorite["data"]["results"]] == ["alpha.md"]
    assert [item["name"] for item in tagged["data"]["results"]] == ["beta.md"]
    assert tagged["data"]["results"][0]["tags"] == ["project"]


def test_recent_search_ignores_paths_outside_notes_directory(tmp_path):
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    note = notes_dir / "inside.md"
    note.write_text("# Inside\nbody", encoding="utf-8")
    outside = tmp_path / "outside.md"
    outside.write_text("# Outside\nbody", encoding="utf-8")
    add_recent(outside, notes_dir=notes_dir)
    add_recent(note, notes_dir=notes_dir)
    client = _client(notes_dir)

    response = client.get(
        "/api/search",
        params={"mode": "recent"},
        headers=_auth(),
    )

    assert response.status_code == 200
    results = response.json()["data"]["results"]
    assert [item["name"] for item in results] == ["inside.md"]


def test_preview_returns_note_content_and_records_recent(tmp_path):
    note = tmp_path / "alpha.md"
    note.write_text("# Alpha\nbody", encoding="utf-8")
    client = _client(tmp_path)

    response = client.get(
        "/api/notes/preview",
        params={"path": str(note)},
        headers=_auth(),
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is True
    assert payload["data"]["content"] == "# Alpha\nbody"
    assert payload["data"]["name"] == "alpha.md"


def test_library_files_list_notes_sources_and_wiki_markdown(tmp_path):
    wiki_dir = tmp_path / "wiki"
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    wiki_dir.mkdir()
    (notes_dir / "alpha.md").write_text("# Alpha\nnote body", encoding="utf-8")
    (notes_dir / "book.pdf").write_bytes(b"%PDF-1.4")
    (notes_dir / "doc.docx").write_bytes(b"docx")
    (notes_dir / "photo.png").write_bytes(b"png")
    (notes_dir / ".assets").mkdir()
    (notes_dir / ".assets" / "clip.png").write_bytes(b"png")
    (wiki_dir / "index.md").write_text("# Wiki Index\n", encoding="utf-8")
    (wiki_dir / "sources").mkdir()
    (wiki_dir / "sources" / "summary_alpha.md").write_text("# Alpha Source\n", encoding="utf-8")
    (wiki_dir / "entities").mkdir()
    (wiki_dir / "entities" / "openai.md").write_text("# OpenAI\nwiki body", encoding="utf-8")
    (wiki_dir / "concepts").mkdir()
    (wiki_dir / "concepts" / "llm.md").write_text("# LLM\nwiki concept", encoding="utf-8")
    client = _client(notes_dir, wiki_dir=wiki_dir)

    source_response = client.get(
        "/api/library/files",
        params={"scope": "notes"},
        headers=_auth(),
    )
    wiki_response = client.get(
        "/api/library/files",
        params={"scope": "wiki"},
        headers=_auth(),
    )

    assert source_response.status_code == 200
    source_items = source_response.json()["data"]["items"]
    assert [item["relative_path"] for item in source_items] == ["book.pdf", "doc.docx", "photo.png"]
    assert [item["kind"] for item in source_items] == ["pdf", "docx", "image"]

    assert wiki_response.status_code == 200
    data = wiki_response.json()["data"]
    assert [group["kind"] for group in data["groups"]] == ["source", "concept", "entity"]
    grouped = {
        group["kind"]: [item["relative_path"] for item in group["items"]]
        for group in data["groups"]
    }
    assert grouped == {
        "source": ["sources/summary_alpha.md"],
        "concept": ["concepts/llm.md"],
        "entity": ["entities/openai.md"],
    }
    assert [item["relative_path"] for item in data["items"]] == [
        "sources/summary_alpha.md",
        "concepts/llm.md",
        "entities/openai.md",
    ]
    assert data["items"][2]["snippet"] == "# OpenAI\nwiki body"


def test_library_preview_reads_notes_and_wiki_but_rejects_outside_paths(tmp_path):
    wiki_dir = tmp_path / "wiki"
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    wiki_dir.mkdir()
    (notes_dir / "photo.png").write_bytes(b"png")
    (notes_dir / "paper.pdf").write_bytes(b"%PDF-1.4")
    (wiki_dir / "concepts").mkdir()
    (wiki_dir / "concepts" / "llm.md").write_text("# LLM\nwiki body", encoding="utf-8")
    outside = tmp_path / "outside.md"
    outside.write_text("secret", encoding="utf-8")
    client = _client(notes_dir, wiki_dir=wiki_dir)

    source_preview = client.get(
        "/api/library/files/preview",
        params={"scope": "notes", "path": "photo.png"},
        headers=_auth(),
    )
    pdf_preview = client.get(
        "/api/library/files/preview",
        params={"scope": "notes", "path": "paper.pdf"},
        headers=_auth(),
    )
    wiki_preview = client.get(
        "/api/library/files/preview",
        params={"scope": "wiki", "path": "concepts/llm.md"},
        headers=_auth(),
    )
    outside_preview = client.get(
        "/api/library/files/preview",
        params={"scope": "wiki", "path": str(outside)},
        headers=_auth(),
    )

    assert source_preview.status_code == 200
    assert source_preview.json()["data"]["render_mode"] == "image"
    assert source_preview.json()["data"]["media_url"].startswith("/api/library/files/raw?")
    assert source_preview.json()["data"]["relative_path"] == "photo.png"
    assert pdf_preview.status_code == 200
    assert pdf_preview.json()["data"]["render_mode"] == "pdf"
    assert pdf_preview.json()["data"]["media_url"].startswith("/api/library/files/raw?")
    assert wiki_preview.status_code == 200
    assert wiki_preview.json()["data"]["content"] == "# LLM\nwiki body"
    assert wiki_preview.json()["data"]["relative_path"] == "concepts/llm.md"
    assert wiki_preview.json()["data"]["render_mode"] == "markdown"
    assert outside_preview.status_code == 400


def test_library_docx_preview_returns_html_with_embedded_images(tmp_path):
    from docx import Document
    from PIL import Image

    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    image_path = tmp_path / "clip.png"
    Image.new("RGB", (2, 2), "blue").save(image_path)
    docx_path = notes_dir / "doc.docx"
    doc = Document()
    doc.add_heading("Doc Title", level=1)
    doc.add_paragraph("First paragraph")
    doc.add_picture(str(image_path))
    doc.save(docx_path)
    client = _client(notes_dir, wiki_dir=tmp_path / "wiki")

    response = client.get(
        "/api/library/files/preview",
        params={"scope": "notes", "path": "doc.docx"},
        headers=_auth(),
    )

    assert response.status_code == 200
    data = response.json()["data"]
    assert data["render_mode"] == "docx_html"
    assert "<h1>Doc Title</h1>" in data["html"]
    assert "<p>First paragraph</p>" in data["html"]
    assert 'src="data:image/png;base64,' in data["html"]


def test_meta_favorite_and_tags_roundtrip(tmp_path):
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")
    client = _client(tmp_path)

    fav = client.post(
        "/api/meta/favorite",
        json={"path": str(note), "favorite": True},
        headers=_auth(),
    )
    tags = client.post(
        "/api/meta/tags",
        json={"path": str(note), "tags": ["AI", "project"]},
        headers=_auth(),
    )
    meta = client.get(
        "/api/meta",
        params={"path": str(note)},
        headers=_auth(),
    )

    assert fav.json()["data"]["favorite"] is True
    assert tags.json()["data"]["tags"] == ["AI", "project"]
    assert meta.json()["data"] == {"favorite": True, "tags": ["AI", "project"]}


def test_web_library_files_support_favorite_and_tags(tmp_path):
    wiki_dir = tmp_path / "wiki"
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    wiki_dir.mkdir()
    source = notes_dir / "paper.pdf"
    source.write_bytes(b"%PDF-1.4")
    (wiki_dir / "concepts").mkdir()
    wiki_page = wiki_dir / "concepts" / "llm.md"
    wiki_page.write_text("# LLM\nwiki body", encoding="utf-8")
    client = _client(notes_dir, wiki_dir=wiki_dir)

    source_fav = client.post(
        "/api/meta/favorite",
        json={"scope": "notes", "path": "paper.pdf", "favorite": True},
        headers=_auth(),
    )
    wiki_fav = client.post(
        "/api/meta/favorite",
        json={"scope": "wiki", "path": "concepts/llm.md", "favorite": True},
        headers=_auth(),
    )
    source_tags = client.post(
        "/api/meta/tags",
        json={"scope": "notes", "path": "paper.pdf", "tags": ["research"]},
        headers=_auth(),
    )
    wiki_tags = client.post(
        "/api/meta/tags",
        json={"scope": "wiki", "path": "concepts/llm.md", "tags": ["research"]},
        headers=_auth(),
    )
    favorite = client.get(
        "/api/search",
        params={"mode": "favorite"},
        headers=_auth(),
    )
    tagged = client.get(
        "/api/search",
        params={"mode": "tag", "q": "research"},
        headers=_auth(),
    )
    wiki_preview = client.get(
        "/api/library/files/preview",
        params={"scope": "wiki", "path": "concepts/llm.md"},
        headers=_auth(),
    )

    assert source_fav.json()["data"]["favorite"] is True
    assert wiki_fav.json()["data"]["favorite"] is True
    assert source_tags.json()["data"]["tags"] == ["research"]
    assert wiki_tags.json()["data"]["tags"] == ["research"]
    favorite_items = favorite.json()["data"]["results"]
    tagged_items = tagged.json()["data"]["results"]
    assert [(item["scope"], item["relative_path"]) for item in favorite_items] == [
        ("notes", "paper.pdf"),
        ("wiki", "concepts/llm.md"),
    ]
    assert [(item["scope"], item["relative_path"]) for item in tagged_items] == [
        ("notes", "paper.pdf"),
        ("wiki", "concepts/llm.md"),
    ]
    assert wiki_preview.json()["data"]["favorite"] is True
    assert wiki_preview.json()["data"]["tags"] == ["research"]


def test_tag_search_without_query_lists_all_tagged_web_files(tmp_path):
    wiki_dir = tmp_path / "wiki"
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    wiki_dir.mkdir()
    note = notes_dir / "alpha.md"
    note.write_text("# Alpha\nbody", encoding="utf-8")
    source = notes_dir / "paper.pdf"
    source.write_bytes(b"%PDF-1.4")
    (wiki_dir / "entities").mkdir()
    wiki_page = wiki_dir / "entities" / "openai.md"
    wiki_page.write_text("# OpenAI\nwiki body", encoding="utf-8")
    client = _client(notes_dir, wiki_dir=wiki_dir)

    client.post(
        "/api/meta/tags",
        json={"scope": "notes", "path": "alpha.md", "tags": ["note"]},
        headers=_auth(),
    )
    client.post(
        "/api/meta/tags",
        json={"scope": "notes", "path": "paper.pdf", "tags": ["source"]},
        headers=_auth(),
    )
    client.post(
        "/api/meta/tags",
        json={"scope": "wiki", "path": "entities/openai.md", "tags": ["wiki"]},
        headers=_auth(),
    )

    response = client.get(
        "/api/search",
        params={"mode": "tag"},
        headers=_auth(),
    )

    assert response.status_code == 200
    results = response.json()["data"]["results"]
    assert [(item["scope"], item["relative_path"], item["tags"]) for item in results] == [
        ("notes", "alpha.md", ["note"]),
        ("notes", "paper.pdf", ["source"]),
        ("wiki", "entities/openai.md", ["wiki"]),
    ]


def test_open_reader_endpoint_invokes_callback_with_resolved_note(tmp_path):
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")
    calls: list[tuple[Path, str]] = []
    client = TestClient(
        create_app(
            notes_dir=tmp_path,
            panel_token="panel-token",
            open_reader=lambda path, query: calls.append((path, query)),
        )
    )

    response = client.post(
        "/api/notes/open-reader",
        json={"path": str(note), "query": "alpha"},
        headers=_auth(),
    )

    assert response.status_code == 200
    assert response.json()["data"] == {"opened": True}
    assert calls == [(note.resolve(), "alpha")]


def test_create_note_endpoint_saves_markdown_to_notes_dir(tmp_path):
    client = _client(tmp_path)

    response = client.post(
        "/api/notes",
        json={"title": "Alpha Note", "content": "first body"},
        headers=_auth(),
    )

    assert response.status_code == 200
    data = response.json()["data"]
    saved = Path(data["path"])
    assert data["name"] == "Alpha_Note.md"
    assert saved.parent == tmp_path
    assert saved.read_text(encoding="utf-8").startswith("---\ntitle: Alpha Note")
    assert "first body" in saved.read_text(encoding="utf-8")


def test_asset_upload_saves_image_and_returns_markdown_snippet(monkeypatch, tmp_path):
    monkeypatch.setattr(api_module, "ocr_image", lambda _img: ["图片文字"], raising=False)
    client = _client(tmp_path)
    image = Image.new("RGB", (2, 2), "blue")
    buf = BytesIO()
    image.save(buf, format="PNG")

    response = client.post(
        "/api/assets",
        files={"file": ("clip.png", buf.getvalue(), "image/png")},
        headers=_auth(),
    )

    assert response.status_code == 200
    data = response.json()["data"]
    assert data["name"].endswith(".png")
    assert data["markdown"].startswith("![](.assets/")
    assert (tmp_path / ".assets" / data["name"]).exists()


def test_upload_preview_save_and_inbox_roundtrip(tmp_path):
    client = _client(tmp_path)
    file_payload = {"file": ("source.md", b"# Source\nbody", "text/markdown")}

    preview = client.post("/api/uploads/preview", files=file_payload, headers=_auth())
    saved = client.post(
        "/api/uploads",
        files={"file": ("source.md", b"# Source\nbody", "text/markdown")},
        headers=_auth(),
    )
    inbox = client.get("/api/inbox", headers=_auth())

    assert preview.status_code == 200
    assert preview.json()["data"]["preview"].startswith("# Source")
    assert saved.status_code == 200
    saved_path = Path(saved.json()["data"]["path"])
    assert saved_path.parent == tmp_path
    assert saved_path.exists()
    assert saved_path.name in [item["name"] for item in inbox.json()["data"]["items"]]

    inbox_preview = client.get(
        "/api/inbox/preview",
        params={"path": str(saved_path)},
        headers=_auth(),
    )
    assert inbox_preview.json()["data"]["content"].startswith("# Source")


def test_delete_inbox_item_removes_note_file(tmp_path):
    client = _client(tmp_path)
    note = tmp_path / "source.md"
    note.write_text("# Source\nbody", encoding="utf-8")

    response = client.delete(
        "/api/inbox",
        params={"path": str(note)},
        headers=_auth(),
    )
    inbox = client.get("/api/inbox", headers=_auth())

    assert response.status_code == 200
    assert response.json()["data"] == {"path": str(note), "name": "source.md"}
    assert not note.exists()
    assert "source.md" not in [item["name"] for item in inbox.json()["data"]["items"]]


def test_delete_inbox_item_rejects_notes_not_in_inbox(tmp_path):
    wiki_dir = tmp_path / "wiki"
    sources_dir = wiki_dir / "sources"
    sources_dir.mkdir(parents=True)
    note = tmp_path / "source.md"
    note.write_text("# Source\nbody", encoding="utf-8")
    (sources_dir / "summary_source.md").write_text("# Source summary\n", encoding="utf-8")
    client = _client(tmp_path, wiki_dir=wiki_dir)

    response = client.delete(
        "/api/inbox",
        params={"path": str(note)},
        headers=_auth(),
    )

    assert response.status_code == 404
    assert note.exists()


def test_query_stream_returns_meta_thinking_chunks_and_done(monkeypatch, tmp_path):
    def fake_query(question, config, *, wiki_dir=None, notes_dir=None, on_meta=None, on_thinking=None):
        assert question == "What is Alpha?"
        assert notes_dir == tmp_path
        if on_meta:
            on_meta(SimpleNamespace(
                question=question,
                answer_type="direct_answer",
                used_pages=["concepts/alpha.md"],
                raw_sources=["alpha.md"],
                suggested_save_title="alpha",
            ))
        if on_thinking:
            on_thinking("checking")
        yield "Alpha "
        yield "answer"

    monkeypatch.setattr(api_module, "query_wiki", fake_query, raising=False)
    monkeypatch.setattr(api_module, "load_llm_config", lambda: object(), raising=False)
    client = _client(tmp_path)

    response = client.post(
        "/api/wiki/query",
        json={"question": "What is Alpha?"},
        headers=_auth(),
    )

    assert response.status_code == 200
    events = [json.loads(line) for line in response.text.splitlines() if line]
    assert [event["type"] for event in events] == ["meta", "thinking", "chunk", "chunk", "done"]
    assert events[0]["meta"]["used_pages"] == ["concepts/alpha.md"]
    assert events[2]["text"] == "Alpha "


def test_query_stream_uses_worker_queue_so_thinking_is_not_blocked():
    source = Path(api_module.__file__).read_text(encoding="utf-8")
    query_route = source.split('@app.post("/api/wiki/query")', 1)[1].split('@app.post("/api/wiki/query/save")', 1)[0]

    assert "threading.Thread" in query_route
    assert ".get(timeout=" in query_route
    assert '"type": "thinking"' in query_route


def test_query_save_endpoint_persists_last_answer(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    target = wiki_dir / "synthesis" / "alpha.md"

    def fake_save(question, answer, used_pages, *, answer_type, raw_sources):
        target.parent.mkdir(parents=True)
        target.write_text(answer, encoding="utf-8")
        return target

    monkeypatch.setattr(api_module, "save_query_answer_as_wiki_page", fake_save, raising=False)
    client = _client(tmp_path, wiki_dir=wiki_dir)

    response = client.post(
        "/api/wiki/query/save",
        json={
            "question": "What is Alpha?",
            "answer": "Alpha answer",
            "used_pages": ["concepts/alpha.md"],
            "raw_sources": ["alpha.md"],
            "answer_type": "direct_answer",
        },
        headers=_auth(),
    )

    assert response.status_code == 200
    assert response.json()["data"]["name"] == "alpha.md"
    assert target.read_text(encoding="utf-8") == "Alpha answer"


def test_graph_endpoint_returns_nodes_edges_and_diagnostics(tmp_path):
    wiki_dir = tmp_path / "wiki"
    (wiki_dir / "entities").mkdir(parents=True)
    (wiki_dir / "concepts").mkdir()
    (wiki_dir / "sources").mkdir()
    (wiki_dir / "index.md").write_text(
        "\n".join([
            "# Wiki Index",
            "## Sources",
            "- [Alpha Source](sources/summary_alpha.md) - source summary",
            "## Entities",
            "- [OpenAI](entities/openai.md) - entity summary",
            "- [Missing](entities/missing.md)",
            "## Concepts",
            "- [LLM](concepts/llm.md)",
        ]),
        encoding="utf-8",
    )
    (wiki_dir / "entities" / "openai.md").write_text(
        "# OpenAI\n\nBuilds models.\n\n## Related\n- [LLM](../concepts/llm.md)\n",
        encoding="utf-8",
    )
    (wiki_dir / "concepts" / "llm.md").write_text(
        "# LLM\n\nA model family.\n\n## Related\n- [OpenAI](../entities/openai.md)\n",
        encoding="utf-8",
    )
    (wiki_dir / "sources" / "summary_alpha.md").write_text("# Alpha Source\n", encoding="utf-8")
    client = _client(tmp_path, wiki_dir=wiki_dir)

    response = client.get("/api/wiki/graph", headers=_auth())

    assert response.status_code == 200
    data = response.json()["data"]
    assert {node["id"] for node in data["nodes"]} == {
        "sources/summary_alpha.md",
        "entities/openai.md",
        "entities/missing.md",
        "concepts/llm.md",
    }
    openai = next(node for node in data["nodes"] if node["id"] == "entities/openai.md")
    missing = next(node for node in data["nodes"] if node["id"] == "entities/missing.md")
    assert openai["degree"] == 2
    assert missing["exists"] is False
    assert missing["id"] in data["diagnostics"]["missing"]
    assert "sources/summary_alpha.md" in data["diagnostics"]["orphan"]
    assert data["edges"] == [{
        "source": "entities/openai.md",
        "target": "concepts/llm.md",
        "kind": "related",
        "bidirectional": True,
    }]


def test_lint_stream_fix_and_rebuild_endpoints(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    finding = LintFinding(
        severity="warn",
        kind="heading_drift",
        location="index.md",
        message="Heading drift",
        suggestion="Use ## Sources",
        priority="P1",
        fixable=True,
    )

    monkeypatch.setattr(api_module, "load_llm_config", lambda: object(), raising=False)
    monkeypatch.setattr(api_module, "lint_wiki", lambda wiki, cfg: iter([finding]), raising=False)
    monkeypatch.setattr(api_module, "save_lint_report", lambda wiki, findings: wiki / "synthesis" / "wiki-lint.md", raising=False)
    monkeypatch.setattr(api_module, "append_wiki_log", lambda *args, **kwargs: None, raising=False)
    monkeypatch.setattr(api_module, "auto_fix", lambda wiki, findings: len(findings), raising=False)
    monkeypatch.setattr(api_module, "static_checks", lambda wiki: [], raising=False)
    monkeypatch.setattr(api_module, "rebuild_index_from_disk", lambda wiki: None, raising=False)
    client = _client(tmp_path, wiki_dir=wiki_dir)

    lint_response = client.post("/api/wiki/lint", headers=_auth())
    events = [json.loads(line) for line in lint_response.text.splitlines() if line]
    assert [event["type"] for event in events] == ["progress", "finding", "done"]
    assert events[1]["finding"]["kind"] == "heading_drift"

    fix_response = client.post(
        "/api/wiki/lint/fix",
        json={"findings": [events[1]["finding"]]},
        headers=_auth(),
    )
    rebuild_response = client.post("/api/wiki/lint/rebuild-index", headers=_auth())

    assert fix_response.json()["data"]["fixed"] == 1
    assert rebuild_response.json()["data"]["findings"] == []


def test_lint_fix_preview_and_apply_endpoints(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    finding = LintFinding(
        severity="warn",
        kind="no_sources",
        location="entities/e.md",
        message="No sources",
        suggestion="Add sources",
        priority="P1",
    )
    preview = api_module.LintFixPreview(files=(
        api_module.LintFixFile(
            path="entities/e.md",
            original="# E\n",
            updated="# E\n\n## Sources\n\n- [A](../sources/a.md)\n",
            issues=("no_sources",),
        ),
    ), summary="Added sources")

    monkeypatch.setattr(api_module, "load_llm_config", lambda: SimpleNamespace(api_key="key"), raising=False)
    monkeypatch.setattr(api_module, "auto_fix", lambda wiki, findings: 0, raising=False)
    monkeypatch.setattr(api_module, "static_checks", lambda wiki: [finding], raising=False)
    monkeypatch.setattr(api_module, "build_llm_fix_preview", lambda wiki, findings, cfg: preview, raising=False)
    monkeypatch.setattr(api_module, "apply_llm_fix_preview", lambda wiki, payload: len(payload.files), raising=False)
    monkeypatch.setattr(api_module, "rebuild_index_from_disk", lambda wiki: None, raising=False)
    client = _client(tmp_path, wiki_dir=wiki_dir)

    preview_response = client.post(
        "/api/wiki/lint/fix-preview",
        json={"findings": [api_module._finding_payload(finding)]},
        headers=_auth(),
    )
    apply_response = client.post(
        "/api/wiki/lint/apply-preview",
        json={"preview": preview_response.json()["data"]["preview"]},
        headers=_auth(),
    )

    assert preview_response.status_code == 200
    assert preview_response.json()["data"]["preview"]["files"][0]["path"] == "entities/e.md"
    assert apply_response.status_code == 200
    assert apply_response.json()["data"]["written"] == 1


def test_lint_fix_preview_sends_only_selected_findings_to_llm(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    selected = [
        LintFinding(
            severity="warn",
            kind="no_sources",
            location=f"entities/selected_{i}.md",
            message="No sources",
            suggestion="Add sources",
            priority="P1",
        )
        for i in range(2)
    ]
    unselected = [
        LintFinding(
            severity="warn",
            kind="orphan",
            location=f"sources/unselected_{i}.md",
            message="No inbound links",
            suggestion="Add a cross-reference",
            priority="P2",
        )
        for i in range(5)
    ]
    captured: dict[str, list[LintFinding]] = {}

    def capture_preview(_wiki, findings, _cfg):
        captured["findings"] = list(findings)
        return api_module.LintFixPreview(files=(), summary="")

    monkeypatch.setattr(api_module, "load_llm_config", lambda: SimpleNamespace(api_key="key"), raising=False)
    monkeypatch.setattr(api_module, "auto_fix", lambda wiki, findings: 0, raising=False)
    monkeypatch.setattr(api_module, "static_checks", lambda wiki: selected + unselected, raising=False)
    monkeypatch.setattr(api_module, "build_llm_fix_preview", capture_preview, raising=False)
    client = _client(tmp_path, wiki_dir=wiki_dir)

    response = client.post(
        "/api/wiki/lint/fix-preview",
        json={"findings": [api_module._finding_payload(finding) for finding in selected]},
        headers=_auth(),
    )

    assert response.status_code == 200
    assert [finding.location for finding in captured["findings"]] == [
        "entities/selected_0.md",
        "entities/selected_1.md",
    ]
    returned_locations = {
        item["location"] for item in response.json()["data"]["findings"]
    }
    assert "sources/unselected_0.md" in returned_locations


def test_lint_fix_preview_returns_json_error_when_llm_preview_fails(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    finding = LintFinding(
        severity="warn",
        kind="no_sources",
        location="entities/e.md",
        message="No sources",
        suggestion="Add sources",
        priority="P1",
    )

    monkeypatch.setattr(api_module, "load_llm_config", lambda: SimpleNamespace(api_key="key"), raising=False)
    monkeypatch.setattr(api_module, "auto_fix", lambda wiki, findings: 0, raising=False)
    monkeypatch.setattr(api_module, "static_checks", lambda wiki: [finding], raising=False)

    def fail_preview(_wiki, _findings, _cfg):
        raise ValueError("LLM returned invalid fix JSON")

    monkeypatch.setattr(api_module, "build_llm_fix_preview", fail_preview, raising=False)
    client = TestClient(
        create_app(notes_dir=tmp_path, panel_token="panel-token", wiki_dir=wiki_dir),
        raise_server_exceptions=False,
    )

    response = client.post(
        "/api/wiki/lint/fix-preview",
        json={"findings": [api_module._finding_payload(finding)]},
        headers=_auth(),
    )

    assert response.status_code == 502
    payload = response.json()
    assert payload["ok"] is False
    assert payload["error"]["code"] == "llm_fix_preview_failed"
    assert "invalid fix JSON" in payload["error"]["message"]


def test_lint_fix_preview_returns_success_when_llm_preview_times_out(monkeypatch, tmp_path):
    wiki_dir = tmp_path / "wiki"
    finding = LintFinding(
        severity="warn",
        kind="no_sources",
        location="entities/e.md",
        message="No sources",
        suggestion="Add sources",
        priority="P1",
    )

    monkeypatch.setattr(api_module, "load_llm_config", lambda: SimpleNamespace(api_key="key"), raising=False)
    monkeypatch.setattr(api_module, "auto_fix", lambda wiki, findings: 0, raising=False)
    monkeypatch.setattr(api_module, "static_checks", lambda wiki: [finding], raising=False)

    def timeout_preview(_wiki, _findings, _cfg):
        raise httpx.ReadTimeout("The read operation timed out")

    monkeypatch.setattr(api_module, "build_llm_fix_preview", timeout_preview, raising=False)
    client = TestClient(
        create_app(notes_dir=tmp_path, panel_token="panel-token", wiki_dir=wiki_dir),
        raise_server_exceptions=False,
    )

    response = client.post(
        "/api/wiki/lint/fix-preview",
        json={"findings": [api_module._finding_payload(finding)]},
        headers=_auth(),
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ok"] is True
    assert payload["data"]["preview"] is None
    assert "超时" in payload["data"]["llm_error"]


def test_ingest_session_rejects_when_llm_is_unconfigured(monkeypatch, tmp_path):
    monkeypatch.setattr(api_module, "app_config", SimpleNamespace(LLM_API_KEY=""), raising=False)
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")
    client = _client(tmp_path)

    response = client.post(
        "/api/wiki/ingest",
        json={"paths": [str(note)]},
        headers=_auth(),
    )

    assert response.status_code == 503
    assert response.json()["error"]["code"] == "llm_unavailable"


def test_ingest_websocket_maps_queue_protocol_to_events(monkeypatch, tmp_path):
    monkeypatch.setattr(api_module, "app_config", SimpleNamespace(LLM_API_KEY="key"), raising=False)
    monkeypatch.setattr(api_module, "load_llm_config", lambda: object(), raising=False)
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")

    def fake_ingest(note_path, config, *, wiki_dir=None, chat_q, user_q):
        chat_q.put("hello")
        chat_q.put("__SELECT_DEFAULT__")
        assert user_q.get(timeout=2) == "默认"
        chat_q.put("plan")
        chat_q.put("__READY__")
        assert user_q.get(timeout=2) == "confirm"
        chat_q.put("__DONE__")
        return wiki_dir / "sources" / "summary_alpha.md"

    client = _client(tmp_path, wiki_dir=tmp_path / "wiki", ingest_runner=fake_ingest)
    session = client.post(
        "/api/wiki/ingest",
        json={"paths": [str(note)]},
        headers=_auth(),
    ).json()["data"]["session_id"]

    with client.websocket_connect(
        f"/api/wiki/ingest/{session}?token=panel-token",
    ) as websocket:
        assert websocket.receive_json()["type"] == "note"
        assert websocket.receive_json() == {"type": "chunk", "text": "hello"}
        assert websocket.receive_json()["type"] == "select"
        websocket.send_json({"type": "input", "text": "默认"})
        assert websocket.receive_json() == {"type": "chunk", "text": "plan"}
        assert websocket.receive_json()["type"] == "ready"
        websocket.send_json({"type": "confirm"})
        assert websocket.receive_json()["type"] == "done"


def test_ingest_websocket_maps_discussion_input_request(monkeypatch, tmp_path):
    monkeypatch.setattr(api_module, "app_config", SimpleNamespace(LLM_API_KEY="key"), raising=False)
    monkeypatch.setattr(api_module, "load_llm_config", lambda: object(), raising=False)
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")

    def fake_ingest(note_path, config, *, wiki_dir=None, chat_q, user_q):
        chat_q.put("Which topic?")
        chat_q.put("__AWAIT_INPUT__")
        assert user_q.get(timeout=2) == "OpenAI"
        chat_q.put("__DONE__")
        return None

    client = _client(tmp_path, wiki_dir=tmp_path / "wiki", ingest_runner=fake_ingest)
    session = client.post(
        "/api/wiki/ingest",
        json={"paths": [str(note)]},
        headers=_auth(),
    ).json()["data"]["session_id"]

    with client.websocket_connect(
        f"/api/wiki/ingest/{session}?token=panel-token",
    ) as websocket:
        assert websocket.receive_json()["type"] == "note"
        assert websocket.receive_json() == {"type": "chunk", "text": "Which topic?"}
        assert websocket.receive_json()["type"] == "input_request"
        websocket.send_json({"type": "input", "text": "OpenAI"})
        assert websocket.receive_json()["type"] == "done"


def test_ingest_websocket_maps_structured_review_events(monkeypatch, tmp_path):
    monkeypatch.setattr(api_module, "app_config", SimpleNamespace(LLM_API_KEY="key"), raising=False)
    monkeypatch.setattr(api_module, "load_llm_config", lambda: object(), raising=False)
    note = tmp_path / "alpha.md"
    note.write_text("body", encoding="utf-8")

    def fake_ingest(note_path, config, *, wiki_dir=None, chat_q, user_q):
        chat_q.put(("__STAGE__", "candidates", "请审阅候选页面"))
        chat_q.put(("__CANDIDATES__", [{
            "path": "entities/openai.md",
            "title": "OpenAI",
            "kind": "entity",
            "selected": True,
            "deep": False,
        }]))
        chat_q.put(("__AWAIT_INPUT__", "candidates", ["candidate_update", "generate_plan"]))
        assert user_q.get(timeout=2) == {
            "type": "candidate_update",
            "path": "entities/openai.md",
            "selected": False,
        }
        chat_q.put(("__PLAN__", [{
            "path": "entities/openai.md",
            "title": "OpenAI",
            "action": "create",
            "reason": "new",
            "contribution": "Builds GPT.",
        }]))
        assert user_q.get(timeout=2) == {
            "type": "command",
            "command": "generate_plan",
        }
        chat_q.put(("__STAGE__", "plan", "请确认写入计划"))
        chat_q.put(("__AWAIT_INPUT__", "plan", ["plan_update", "execute"]))
        assert user_q.get(timeout=2) == {
            "type": "plan_update",
            "path": "entities/openai.md",
            "action": "skip",
        }
        chat_q.put("__DONE__")
        return None

    client = _client(tmp_path, wiki_dir=tmp_path / "wiki", ingest_runner=fake_ingest)
    session = client.post(
        "/api/wiki/ingest",
        json={"paths": [str(note)]},
        headers=_auth(),
    ).json()["data"]["session_id"]

    with client.websocket_connect(
        f"/api/wiki/ingest/{session}?token=panel-token",
    ) as websocket:
        assert websocket.receive_json()["type"] == "note"
        assert websocket.receive_json() == {
            "type": "stage",
            "stage": "candidates",
            "status": "请审阅候选页面",
        }
        candidates = websocket.receive_json()
        assert candidates["type"] == "candidates"
        assert candidates["candidates"][0]["path"] == "entities/openai.md"
        input_request = websocket.receive_json()
        assert input_request == {
            "type": "input_request",
            "mode": "candidates",
            "actions": ["candidate_update", "generate_plan"],
        }
        websocket.send_json({
            "type": "candidate_update",
            "path": "entities/openai.md",
            "selected": False,
        })
        assert websocket.receive_json()["type"] == "plan"
        websocket.send_json({"type": "command", "command": "generate_plan"})
        assert websocket.receive_json()["type"] == "stage"
        assert websocket.receive_json()["mode"] == "plan"
        websocket.send_json({
            "type": "plan_update",
            "path": "entities/openai.md",
            "action": "skip",
        })
        assert websocket.receive_json()["type"] == "done"
