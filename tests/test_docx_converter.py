import pytest
from pathlib import Path
from docx import Document
from PIL import Image

from converter.docx_converter import docx_to_markdown
from converter.ocr_converter import OCRUnavailableError


@pytest.fixture
def sample_docx(tmp_path):
    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


def test_extracts_heading(sample_docx):
    result = docx_to_markdown(sample_docx)
    assert "# Test Heading" in result


def test_extracts_paragraphs(sample_docx):
    result = docx_to_markdown(sample_docx)
    assert "First paragraph." in result
    assert "Second paragraph." in result


def test_returns_string(sample_docx):
    assert isinstance(docx_to_markdown(sample_docx), str)


def test_missing_file_raises():
    with pytest.raises(FileNotFoundError):
        docx_to_markdown(Path("/nonexistent/file.docx"))


def test_extracts_inline_picture_ocr(tmp_path, monkeypatch):
    from converter import docx_converter

    image = tmp_path / "shot.png"
    Image.new("RGB", (10, 10), "white").save(image)

    doc = Document()
    doc.add_paragraph("Before image")
    doc.add_picture(str(image))
    path = tmp_path / "image.docx"
    doc.save(str(path))

    monkeypatch.setattr(docx_converter, "ocr_image", lambda _img: ["图片里的文字", "第二行"])

    result = docx_to_markdown(path)

    assert "Before image" in result
    assert "> [图片文字] 图片里的文字" in result
    assert "> 第二行" in result


def test_inline_picture_ocr_unavailable_keeps_text(tmp_path, monkeypatch):
    from converter import docx_converter

    image = tmp_path / "shot.png"
    Image.new("RGB", (10, 10), "white").save(image)

    doc = Document()
    doc.add_paragraph("Text survives")
    doc.add_picture(str(image))
    path = tmp_path / "image.docx"
    doc.save(str(path))

    def fail(_img):
        raise OCRUnavailableError("missing")

    monkeypatch.setattr(docx_converter, "ocr_image", fail)

    result = docx_to_markdown(path)

    assert "Text survives" in result
    assert "[图片文字]" not in result
